import subprocess
import sys
from io import BytesIO
from pathlib import Path

import httpx
import pytest
from infinity_context_server_harness import run_infinity_context_server
from resource_guards import docling_worker_timeout_seconds, skip_if_low_temp_space

pytest.importorskip("docling")
pytest.importorskip("docx")


def test_docling_asset_extraction_persists_structured_artifacts(tmp_path: Path) -> None:
    skip_if_low_temp_space(tmp_path, label="Docling asset e2e")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with (
        run_infinity_context_server(
            tmp_path,
            database_name="docling-asset-extraction.db",
            extra_env={
                "MEMORY_ASSET_STORAGE_DIR": str(tmp_path / "assets"),
                "MEMORY_EXTRACTION_DEFAULT_PROFILE": "standard_docling",
            },
        ) as server,
        httpx.Client(
            base_url=server.base_url,
            headers={"Authorization": f"Bearer {server.token}"},
            timeout=10,
        ) as client,
    ):
        upload = client.post(
            "/v1/assets",
            params={
                "space_slug": "docling-e2e",
                "memory_scope_external_ref": "frontend",
                "thread_external_ref": "quick-capture",
                "filename": "docling-memory.docx",
                "extract": "true",
            },
            content=_sample_docx_bytes(),
            headers={"Content-Type": content_type},
        )
        assert upload.status_code == 201, upload.text
        extraction_id = upload.json()["data"]["extraction"]["id"]

        worker = subprocess.run(
            [
                sys.executable,
                "-m",
                "infinity_context_server.worker",
                "--once",
                "--limit",
                "10",
            ],
            env=server.env,
            text=True,
            capture_output=True,
            timeout=docling_worker_timeout_seconds(server.env),
            check=False,
        )
        assert worker.returncode == 0, worker.stdout + worker.stderr

        extracted = client.get(f"/v1/asset-extractions/{extraction_id}")
        assert extracted.status_code == 200, extracted.text
        extraction = extracted.json()["data"]
        assert extraction["status"] == "succeeded"
        assert extraction["parser_name"] == "docling_document"
        assert len(extraction["result_document_ids"]) == 1
        assert extraction["metadata"]["docling_table_count"] >= 1

        artifact_types = {item["artifact_type"] for item in extraction["artifacts"]}
        assert artifact_types == {
            "extracted_json",
            "markdown",
            "media_manifest",
            "normalized_json",
            "table_html",
        }

        normalized = next(
            item for item in extraction["artifacts"] if item["artifact_type"] == "normalized_json"
        )
        normalized_download = client.get(f"/v1/extraction-artifacts/{normalized['id']}/download")
        assert normalized_download.status_code == 200, normalized_download.text
        assert normalized_download.headers["content-type"].startswith("application/json")
        assert b"schema_name" in normalized_download.content

        table = next(
            item for item in extraction["artifacts"] if item["artifact_type"] == "table_html"
        )
        table_download = client.get(f"/v1/extraction-artifacts/{table['id']}/download")
        assert table_download.status_code == 200, table_download.text
        assert table_download.headers["content-type"].startswith("text/html")
        assert b"<table" in table_download.content

        chunks = client.get(f"/v1/documents/{extraction['result_document_ids'][0]}/chunks")
        assert chunks.status_code == 200, chunks.text
        chunk_text = " ".join(item["text"] for item in chunks.json()["data"])
        assert "Docling API memory scope evidence" in chunk_text
        assert "frontend" in chunk_text


def _sample_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Docling API memory scope evidence", 1)
    document.add_paragraph("Alex approved frontend quick capture.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Scope"
    table.cell(0, 1).text = "Thread"
    table.cell(1, 0).text = "frontend"
    table.cell(1, 1).text = "quick-capture"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
