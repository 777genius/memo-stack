import shutil
import subprocess
import sys
import wave
from io import BytesIO
from pathlib import Path

import httpx
import pytest
from infinity_context_server_harness import run_infinity_context_server
from PIL import Image

pytest.importorskip("docling")
pytest.importorskip("docx")


@pytest.mark.skipif(shutil.which("ffmpeg") is None, reason="ffmpeg is not installed")
@pytest.mark.skipif(shutil.which("ffprobe") is None, reason="ffprobe is not installed")
def test_multimodal_asset_memory_acceptance_flow(tmp_path: Path) -> None:
    with (
        run_infinity_context_server(
            tmp_path,
            database_name="multimodal-acceptance.db",
            extra_env={
                "MEMORY_ASSET_STORAGE_DIR": str(tmp_path / "assets"),
                "MEMORY_EXTRACTION_DEFAULT_PROFILE": "standard_docling",
                "MEMORY_TRANSCRIPTION_PROVIDER": "disabled",
            },
        ) as server,
        httpx.Client(
            base_url=server.base_url,
            headers={"Authorization": f"Bearer {server.token}"},
            timeout=20,
        ) as client,
    ):
        fact = client.post(
            "/v1/facts",
            json={
                "space_slug": "multimodal-acceptance",
                "memory_scope_external_ref": "frontend",
                "thread_external_ref": "review",
                "text": (
                    "Alex approved multimodal memory acceptance for pdf, docx, image, "
                    "audio and video assets."
                ),
                "kind": "note",
                "source_refs": [{"source_type": "manual", "source_id": "acceptance"}],
            },
            headers={"Idempotency-Key": "multimodal-acceptance-fact"},
        )
        assert fact.status_code == 201, fact.text

        cases = (
            _FixtureCase(
                filename="acceptance.pdf",
                content_type="application/pdf",
                content=_sample_pdf_bytes("Alex multimodal pdf acceptance evidence"),
                expected_artifacts={"extracted_json", "markdown"},
                expected_text="pdf acceptance evidence",
                expected_evidence_modalities={"document"},
            ),
            _FixtureCase(
                filename="acceptance.docx",
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=_sample_docx_bytes(),
                expected_artifacts={"extracted_json", "markdown", "normalized_json"},
                expected_text="Docx acceptance evidence",
                expected_evidence_modalities={"document"},
            ),
            _FixtureCase(
                filename="acceptance.png",
                content_type="image/png",
                content=_sample_png_bytes(),
                expected_artifacts={"extracted_json", "image_regions", "markdown"},
                expected_text="Image asset evidence",
                expected_evidence_modalities={"image"},
                expected_reason_codes={"visual_text_match"},
            ),
            _FixtureCase(
                filename="acceptance.wav",
                content_type="audio/wav",
                content=_sample_wav_bytes(),
                expected_artifacts={"extracted_json", "markdown", "media_manifest"},
                expected_text="Media asset evidence",
                expected_evidence_modalities={"audio"},
                expected_reason_codes={"audio_evidence_match"},
            ),
            _FixtureCase(
                filename="acceptance.mp4",
                content_type="video/mp4",
                content=_sample_mp4_bytes(tmp_path),
                expected_artifacts={"extracted_json", "markdown", "media_manifest", "keyframe"},
                expected_text="Media asset evidence",
                expected_evidence_modalities={"image", "video", "time_range"},
                expected_reason_codes={"keyframe_match"},
            ),
        )

        extraction_ids = []
        for case in cases:
            upload = client.post(
                "/v1/assets",
                params={
                    "space_slug": "multimodal-acceptance",
                    "memory_scope_external_ref": "frontend",
                    "thread_external_ref": "review",
                    "filename": case.filename,
                    "extract": "true",
                    "estimated_media_seconds": 1,
                },
                content=case.content,
                headers={"Content-Type": case.content_type},
            )
            assert upload.status_code == 201, upload.text
            extraction_ids.append((case, upload.json()["data"]["extraction"]["id"]))

        worker = subprocess.run(
            [
                sys.executable,
                "-m",
                "infinity_context_server.worker",
                "--once",
                "--limit",
                "20",
            ],
            env=server.env,
            text=True,
            capture_output=True,
            timeout=180,
            check=False,
        )
        assert worker.returncode == 0, worker.stdout + worker.stderr

        for case, extraction_id in extraction_ids:
            extracted = client.get(f"/v1/asset-extractions/{extraction_id}")
            assert extracted.status_code == 200, extracted.text
            extraction = extracted.json()["data"]
            assert extraction["status"] == "succeeded"
            assert case.expected_artifacts.issubset(
                {item["artifact_type"] for item in extraction["artifacts"]}
            )
            assert extraction["result_document_ids"], case.filename

            chunks = client.get(f"/v1/documents/{extraction['result_document_ids'][0]}/chunks")
            assert chunks.status_code == 200, chunks.text
            chunk_text = " ".join(item["text"] for item in chunks.json()["data"])
            assert case.expected_text in chunk_text

            suggestions = client.post(
                "/v1/link-suggestions",
                json={
                    "space_slug": "multimodal-acceptance",
                    "memory_scope_external_ref": "frontend",
                    "thread_external_ref": "review",
                    "source_type": "document",
                    "source_id": extraction["result_document_ids"][0],
                    "text": f"Alex multimodal acceptance {case.filename} {chunk_text[:500]}",
                    "persist": True,
                    "limit": 8,
                },
            )
            assert suggestions.status_code == 200, suggestions.text
            fact_candidate = next(
                item
                for item in suggestions.json()["data"]["candidates"]
                if item["target_type"] == "fact"
            )
            chunk_candidate = next(
                item
                for item in suggestions.json()["data"]["candidates"]
                if item["target_type"] == "chunk"
                and item["metadata"].get("document_id") == extraction["result_document_ids"][0]
            )
            chunk_metadata = chunk_candidate["metadata"]
            assert case.expected_evidence_modalities.issubset(
                set(chunk_metadata["evidence_modalities"])
            )
            assert case.expected_reason_codes.issubset(set(chunk_metadata["reason_codes"]))
            if "image" in case.expected_evidence_modalities:
                assert chunk_metadata["evidence_has_bbox_ref"] is True
            if "time_range" in case.expected_evidence_modalities:
                assert chunk_metadata["evidence_has_time_range_ref"] is True
            approved = client.post(
                f"/v1/context-link-suggestions/{fact_candidate['suggestion_id']}/review",
                json={"action": "approve", "reason": f"accepted {case.filename} evidence"},
            )
            assert approved.status_code == 200, approved.text
            assert approved.json()["data"]["link"]["target_type"] == "fact"

        anchors = client.post(
            "/v1/anchors/backfill",
            json={
                "space_slug": "multimodal-acceptance",
                "memory_scope_external_ref": "frontend",
                "limit_per_source": 50,
            },
        )
        assert anchors.status_code == 200, anchors.text
        anchor_keys = {
            (item["kind"], item["normalized_key"])
            for item in anchors.json()["data"]["anchors"]
        }
        assert ("person", "alex") in anchor_keys


class _FixtureCase:
    def __init__(
        self,
        *,
        filename: str,
        content_type: str,
        content: bytes,
        expected_artifacts: set[str],
        expected_text: str,
        expected_evidence_modalities: set[str],
        expected_reason_codes: set[str] | None = None,
    ) -> None:
        self.filename = filename
        self.content_type = content_type
        self.content = content
        self.expected_artifacts = expected_artifacts
        self.expected_text = expected_text
        self.expected_evidence_modalities = expected_evidence_modalities
        self.expected_reason_codes = expected_reason_codes or {"text_match"}


def _sample_pdf_bytes(text: str) -> bytes:
    stream = f"BT /F1 18 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    return (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        + f"5 0 obj << /Length {len(stream)} >> stream\n".encode("ascii")
        + stream
        + b"\nendstream endobj\nxref\n0 6\n0000000000 65535 f \n"
        b"0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n"
        b"0000000241 00000 n \n0000000311 00000 n \n"
        b"trailer << /Root 1 0 R /Size 6 >>\nstartxref\n449\n%%EOF\n"
    )


def _sample_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Docx acceptance evidence", 1)
    document.add_paragraph("Alex approved docx acceptance evidence for memory.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _sample_png_bytes() -> bytes:
    image = Image.new("RGB", (120, 40), color=(255, 255, 255))
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _sample_wav_bytes() -> bytes:
    buffer = BytesIO()
    with wave.open(buffer, "wb") as wav:
        wav.setnchannels(1)
        wav.setsampwidth(2)
        wav.setframerate(8000)
        wav.writeframes(b"\x00\x00" * 8000)
    return buffer.getvalue()


def _sample_mp4_bytes(tmp_path: Path) -> bytes:
    video_path = tmp_path / "acceptance.mp4"
    subprocess.run(
        [
            "ffmpeg",
            "-y",
            "-v",
            "error",
            "-f",
            "lavfi",
            "-i",
            "color=c=black:s=32x32:d=1",
            "-pix_fmt",
            "yuv420p",
            str(video_path),
        ],
        check=True,
        capture_output=True,
        timeout=20,
    )
    return video_path.read_bytes()
